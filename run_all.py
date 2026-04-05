import os
import json
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import roc_curve, auc, precision_recall_curve, confusion_matrix
from joblib import dump
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

from src.config import CFG
from src.ncbi_fetch import download_h5n5_and_negative
from src.preprocess import filter_sequences
from src.features import build_kmer_features
from src.train_classical_svm import train_classical_svm
from src.train_quantum_svm import train_quantum_svm
from src.train_random_forest import train_random_forest
from src.train_naive_bayes import train_naive_bayes
from src.train_decision_tree import train_decision_tree
from src.evaluate import evaluate

# Set plotting style
sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (12, 6)

class PimaEDA:
    """Exploratory Data Analysis for Pima Indians Diabetes Dataset"""
    
    def __init__(self, df):
        self.df = df
        self.figures = []
    
    def load_pima_data(self, csv_path="data/raw/pima_indians_diabetes.csv"):
        """Load Pima dataset or download if not exists"""
        if not os.path.exists(csv_path):
            print("Downloading Pima Indians Diabetes dataset...")
            os.makedirs(os.path.dirname(csv_path), exist_ok=True)
            # Using seaborn to load
            try:
                self.df = pd.read_csv("https://raw.githubusercontent.com/jbrownlee/Datasets/master/pima-indians-diabetes.csv")
                cols = ['Pregnancies', 'Glucose', 'BloodPressure', 'SkinThickness', 
                       'Insulin', 'BMI', 'DiabetesPedigreeFunction', 'Age', 'Outcome']
                self.df.columns = cols
                self.df.to_csv(csv_path, index=False)
            except:
                raise ValueError("Could not download Pima dataset. Please provide CSV manually.")
        else:
            self.df = pd.read_csv(csv_path)
        return self.df
    
    def basic_statistics(self):
        """Generate basic statistics"""
        stats = {
            'shape': self.df.shape,
            'missing_values': self.df.isnull().sum().to_dict(),
            'dtypes': self.df.dtypes.to_dict(),
            'describe': self.df.describe().to_dict()
        }
        return stats
    
    def distribution_plots(self):
        """Create distribution plots for all features"""
        fig, axes = plt.subplots(3, 3, figsize=(15, 12))
        axes = axes.flatten()
        
        for idx, col in enumerate(self.df.columns):
            axes[idx].hist(self.df[col], bins=30, edgecolor='black', alpha=0.7, color='skyblue')
            axes[idx].set_title(f'Distribution of {col}', fontsize=10, fontweight='bold')
            axes[idx].set_xlabel(col)
            axes[idx].set_ylabel('Frequency')
        
        plt.tight_layout()
        plt.savefig('data/processed/pima_distributions.png', dpi=300, bbox_inches='tight')
        self.figures.append(('distributions', 'data/processed/pima_distributions.png'))
        plt.close()
        print("✓ Distribution plots saved")
    
    def correlation_heatmap(self):
        """Create correlation matrix heatmap"""
        fig, ax = plt.subplots(figsize=(10, 8))
        corr = self.df.corr()
        sns.heatmap(corr, annot=True, fmt='.2f', cmap='coolwarm', center=0, 
                   square=True, ax=ax, cbar_kws={'label': 'Correlation'})
        ax.set_title('Feature Correlation Matrix - Pima Diabetes', fontsize=14, fontweight='bold')
        plt.tight_layout()
        plt.savefig('data/processed/pima_correlation.png', dpi=300, bbox_inches='tight')
        self.figures.append(('correlation', 'data/processed/pima_correlation.png'))
        plt.close()
        print("✓ Correlation heatmap saved")
    
    def class_distribution(self):
        """Visualize class balance"""
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4))
        
        # Count plot
        outcome_counts = self.df['Outcome'].value_counts()
        colors = ['#2ecc71', '#e74c3c']
        ax1.bar(['No Diabetes (0)', 'Diabetes (1)'], outcome_counts.values, color=colors, edgecolor='black')
        ax1.set_ylabel('Count', fontsize=11)
        ax1.set_title('Class Distribution', fontsize=12, fontweight='bold')
        for i, v in enumerate(outcome_counts.values):
            ax1.text(i, v + 5, str(v), ha='center', fontweight='bold')
        
        # Pie chart
        ax2.pie(outcome_counts.values, labels=['No Diabetes', 'Diabetes'], 
               autopct='%1.1f%%', colors=colors, startangle=90)
        ax2.set_title('Class Balance', fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig('data/processed/pima_class_distribution.png', dpi=300, bbox_inches='tight')
        self.figures.append(('class_distribution', 'data/processed/pima_class_distribution.png'))
        plt.close()
        print("✓ Class distribution plot saved")
    
    def feature_vs_target(self):
        """Box plots of features by class"""
        fig, axes = plt.subplots(3, 3, figsize=(15, 12))
        axes = axes.flatten()
        
        feature_cols = [col for col in self.df.columns if col != 'Outcome']
        
        for idx, col in enumerate(feature_cols):
            self.df.boxplot(column=col, by='Outcome', ax=axes[idx])
            axes[idx].set_title(f'{col} by Outcome', fontsize=10, fontweight='bold')
            axes[idx].set_xlabel('Outcome (0=No, 1=Yes)')
            axes[idx].set_ylabel(col)
        
        plt.suptitle('Feature Distribution by Class', fontsize=14, fontweight='bold', y=1.00)
        plt.tight_layout()
        plt.savefig('data/processed/pima_feature_vs_target.png', dpi=300, bbox_inches='tight')
        self.figures.append(('feature_vs_target', 'data/processed/pima_feature_vs_target.png'))
        plt.close()
        print("✓ Feature vs target plots saved")


class ResultsVisualizer:
    """Visualize ML results"""
    
    @staticmethod
    def plot_model_comparison(classical_acc, quantum_acc, rf_acc, nb_acc, dt_acc, 
                             pima_c_acc=None, pima_q_acc=None, pima_rf_acc=None, pima_nb_acc=None, pima_dt_acc=None):
        """Compare model accuracies"""
        fig, ax = plt.subplots(figsize=(14, 8))
        
        models = []
        accs = []
        colors = []
        
        models.extend(['Classical SVM\n(H5N5)', 'Quantum SVM\n(H5N5)', 'Random Forest\n(H5N5)', 'Naive Bayes\n(H5N5)', 'Decision Tree\n(H5N5)'])
        accs.extend([classical_acc, quantum_acc, rf_acc, nb_acc, dt_acc])
        colors.extend(['#3498db', '#9b59b6', '#e74c3c', '#f1c40f', '#e67e22'])
        
        if pima_c_acc and pima_q_acc and pima_rf_acc and pima_nb_acc and pima_dt_acc:
            models.extend(['Classical SVM\n(Pima)', 'Quantum SVM\n(Pima)', 'Random Forest\n(Pima)', 'Naive Bayes\n(Pima)', 'Decision Tree\n(Pima)'])
            accs.extend([pima_c_acc, pima_q_acc, pima_rf_acc, pima_nb_acc, pima_dt_acc])
            colors.extend(['#2980b9', '#8e44ad', '#c0392b', '#d35400', '#f39c12'])
        
        bars = ax.bar(models, accs, color=colors, edgecolor='black', linewidth=1.5, alpha=0.85)
        ax.set_ylabel('Accuracy', fontsize=12, fontweight='bold')
        ax.set_title('Model Performance Comparison: H5N5 & Pima Diabetes', fontsize=15, fontweight='bold')
        ax.set_ylim([0, 1.05])
        
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                   f'{height*100:.2f}%', ha='center', va='bottom', fontweight='bold', fontsize=10)
        
        plt.tight_layout()
        plt.savefig('data/processed/model_comparison.png', dpi=300, bbox_inches='tight')
        plt.close()
        print("✓ Model comparison plot saved")
    
    @staticmethod
    def plot_confusion_matrices(y_test, predictions, title_suffix=""):
        """Plot confusion matrices for all models side-by-side"""
        n_models = len(predictions)
        fig, axes = plt.subplots(1, n_models, figsize=(6 * n_models, 5))
        
        if n_models == 1:
            axes = [axes]
            
        colors = ['Blues', 'Purples', 'Greens', 'YlOrBr', 'Oranges']
        
        for idx, (name, y_pred) in enumerate(predictions.items()):
            cm = confusion_matrix(y_test, y_pred)
            sns.heatmap(cm, annot=True, fmt='d', cmap=colors[idx % len(colors)], ax=axes[idx], cbar=False)
            axes[idx].set_title(f'{name} {title_suffix}', fontsize=12, fontweight='bold')
            axes[idx].set_ylabel('True Label')
            axes[idx].set_xlabel('Predicted Label')
        
        plt.tight_layout()
        plt.savefig(f'data/processed/confusion_matrices{title_suffix.replace(" ", "_")}.png', dpi=300, bbox_inches='tight')
        plt.close()


def train_pima_models(X_train, X_test, y_train, y_test):
    """Train all models on Pima dataset"""
    print("\n" + "="*70)
    print("PIMA DIABETES CLASSIFICATION")
    print("="*70)
    
    # Scale features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    # Quantum features (reduced)
    pca = None
    if X_train_scaled.shape[1] > CFG.n_qubits:
        from sklearn.decomposition import PCA
        pca = PCA(n_components=CFG.n_qubits)
        X_train_q = pca.fit_transform(X_train_scaled)
        X_test_q = pca.transform(X_test_scaled)
        print(f"Reduced features to {CFG.n_qubits} dimensions for Quantum SVM using PCA")
    else:
        X_train_q = X_train_scaled
        X_test_q = X_test_scaled
    
    # Classical SVM
    print("\nTraining Classical SVM on Pima...")
    pima_classical = train_classical_svm(X_train_scaled, y_train)
    acc_pima_c = evaluate(pima_classical.model, X_test_scaled, y_test, "Classical SVM (Pima)")
    
    # Quantum SVM
    print("\nTraining Quantum SVM on Pima...")
    pima_quantum = train_quantum_svm(X_train_q, y_train, n_qubits=CFG.n_qubits)
    acc_pima_q = evaluate(pima_quantum.model, X_test_q, y_test, "Quantum SVM (Pima)")

    # Random Forest
    print("\nTraining Random Forest on Pima...")
    pima_rf = train_random_forest(X_train_scaled, y_train)
    acc_pima_rf = evaluate(pima_rf.model, X_test_scaled, y_test, "Random Forest (Pima)")
    
    # Naive Bayes
    print("\nTraining Naive Bayes on Pima...")
    pima_nb = train_naive_bayes(X_train_scaled, y_train)
    acc_pima_nb = evaluate(pima_nb.model, X_test_scaled, y_test, "Naive Bayes (Pima)")
    
    # Decision Tree
    print("\nTraining Decision Tree on Pima...")
    pima_dt = train_decision_tree(X_train_scaled, y_train)
    acc_pima_dt = evaluate(pima_dt.model, X_test_scaled, y_test, "Decision Tree (Pima)")
    
    dump(pima_classical.model, "data/processed/pima_classical_svm.joblib")
    dump(pima_quantum.model, "data/processed/pima_quantum_svm.joblib")
    dump(pima_rf.model, "data/processed/pima_rf.joblib")
    dump(pima_nb.model, "data/processed/pima_nb.joblib")
    dump(pima_dt.model, "data/processed/pima_dt.joblib")
    
    return acc_pima_c, acc_pima_q, acc_pima_rf, acc_pima_nb, acc_pima_dt, X_test_scaled, X_test_q, y_test, pima_classical, pima_quantum, pima_rf, pima_nb, pima_dt


def generate_word_report(h5n5_results, pima_results, pima_eda_stats):
    """Generate comprehensive Word report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ML Classification Report', 0)
    title.alignment = 1  # Center
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').alignment = 1
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive analysis of two classification tasks using '
        'Classical and Quantum Support Vector Machines (SVMs). The analysis includes the H5N5 '
        'influenza virus classification task and the Pima Indians Diabetes prediction task, '
        'with detailed exploratory data analysis, model performance metrics, and visualizations.'
    )
    
    # H5N5 Results
    doc.add_heading('1. H5N5 Influenza Classification', level=1)
    doc.add_heading('Dataset Overview', level=2)
    doc.add_paragraph(f"Total Samples: {h5n5_results['n_samples']}")
    doc.add_paragraph(f"Positive Samples (H5N5): {h5n5_results['n_pos']}")
    doc.add_paragraph(f"Negative Samples (Other H5Nx): {h5n5_results['n_neg']}")
    doc.add_paragraph(f"Feature Type: {h5n5_results['kmer_k']}-mer k-mers")
    doc.add_paragraph(f"Feature Dimensions (after SVD): {h5n5_results['svd_components']}")
    
    doc.add_heading('Results', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Accuracy'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Classical SVM (RBF)'
    row_cells[1].text = f"{h5n5_results['classical_acc']*100:.2f}%"
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Quantum SVM (QSVC)'
    row_cells[1].text = f"{h5n5_results['quantum_acc']*100:.2f}%"

    row_cells = table.rows[3].cells
    row_cells[0].text = 'Random Forest'
    row_cells[1].text = f"{h5n5_results['rf_acc']*100:.2f}%"
    
    row_cells = table.rows[4].cells
    row_cells[0].text = 'Naive Bayes'
    row_cells[1].text = f"{h5n5_results['nb_acc']*100:.2f}%"
    
    row_cells = table.rows[5].cells
    row_cells[0].text = 'Decision Tree'
    row_cells[1].text = f"{h5n5_results['dt_acc']*100:.2f}%"
    
    # Pima Results
    doc.add_heading('2. Pima Indians Diabetes Classification', level=1)
    doc.add_heading('Exploratory Data Analysis', level=2)
    
    if pima_eda_stats:
        doc.add_paragraph(f"Dataset Shape: {pima_eda_stats['shape']}")
        doc.add_paragraph(f"Missing Values: {pima_eda_stats['missing_values']}")
    
    # Add EDA figures
    if os.path.exists('data/processed/pima_distributions.png'):
        doc.add_heading('Feature Distributions', level=3)
        doc.add_picture('data/processed/pima_distributions.png', width=Inches(6))
    
    if os.path.exists('data/processed/pima_correlation.png'):
        doc.add_heading('Feature Correlations', level=3)
        doc.add_picture('data/processed/pima_correlation.png', width=Inches(5.5))
    
    if os.path.exists('data/processed/pima_class_distribution.png'):
        doc.add_heading('Class Distribution', level=3)
        doc.add_picture('data/processed/pima_class_distribution.png', width=Inches(6))
    
    if os.path.exists('data/processed/pima_feature_vs_target.png'):
        doc.add_heading('Feature Analysis by Class', level=3)
        doc.add_picture('data/processed/pima_feature_vs_target.png', width=Inches(6))
    
    doc.add_heading('Classification Results', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Accuracy'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Classical SVM'
    row_cells[1].text = f"{pima_results['classical_acc']*100:.2f}%"
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Quantum SVM'
    row_cells[1].text = f"{pima_results['quantum_acc']*100:.2f}%"

    row_cells = table.rows[3].cells
    row_cells[0].text = 'Random Forest'
    row_cells[1].text = f"{pima_results['rf_acc']*100:.2f}%"
    
    row_cells = table.rows[4].cells
    row_cells[0].text = 'Naive Bayes'
    row_cells[1].text = f"{pima_results['nb_acc']*100:.2f}%"
    
    row_cells = table.rows[5].cells
    row_cells[0].text = 'Decision Tree'
    row_cells[1].text = f"{pima_results['dt_acc']*100:.2f}%"
    
    # Model Comparison
    doc.add_heading('3. Model Comparison', level=1)
    if os.path.exists('data/processed/model_comparison.png'):
        doc.add_picture('data/processed/model_comparison.png', width=Inches(6))
    
    # Confusion Matrices
    doc.add_heading('4. Detailed Metrics', level=1)
    if os.path.exists('data/processed/confusion_matrices_H5N5.png'):
        doc.add_heading('H5N5 Confusion Matrices', level=2)
        doc.add_picture('data/processed/confusion_matrices_H5N5.png', width=Inches(6))
    
    if os.path.exists('data/processed/confusion_matrices_Pima.png'):
        doc.add_heading('Pima Confusion Matrices', level=2)
        doc.add_picture('data/processed/confusion_matrices_Pima.png', width=Inches(6))
    
    # Conclusions
    doc.add_heading('5. Conclusions', level=1)
    doc.add_paragraph(
        'This analysis demonstrates the application of both classical and quantum machine learning '
        'approaches to real-world classification problems. The results show comparative performance '
        'between traditional RBF SVM and quantum-enhanced QSVC models.'
    )
    
    # Save document
    doc.save('data/processed/ML_Classification_Report.docx')
    print("\n✓ Report saved: data/processed/ML_Classification_Report.docx")


def main():
    os.makedirs("data/raw", exist_ok=True)
    os.makedirs("data/processed", exist_ok=True)

    # ==================== H5N5 ANALYSIS ====================
    print("\n" + "="*70)
    print("H5N5 INFLUENZA CLASSIFICATION")
    print("="*70)
    
    print("Downloading sequences from NCBI...")
    pos, neg = download_h5n5_and_negative(
        email=CFG.ncbi_email,
        api_key=CFG.ncbi_api_key,
        segment_keyword=CFG.target_segment_keyword,
        h5n5_max=CFG.h5n5_max,
        negative_max=CFG.negative_max
    )

    pos_f = filter_sequences(pos, CFG.min_seq_len, CFG.max_ambiguous_frac)
    neg_f = filter_sequences(neg, CFG.min_seq_len, CFG.max_ambiguous_frac)

    print(f"Filtered H5N5: {len(pos_f)}")
    print(f"Filtered NEG : {len(neg_f)}")

    n = min(len(pos_f), len(neg_f))
    pos_f = pos_f[:n]
    neg_f = neg_f[:n]

    rows = []
    for acc, seq in pos_f:
        rows.append({"accession": acc, "sequence": seq, "label": 1})
    for acc, seq in neg_f:
        rows.append({"accession": acc, "sequence": seq, "label": 0})

    df_h5n5 = pd.DataFrame(rows).sample(frac=1.0, random_state=CFG.random_state).reset_index(drop=True)
    df_h5n5.to_csv("data/processed/h5n5_dataset.csv", index=False)
    print("Saved: data/processed/h5n5_dataset.csv")

    sequences = df_h5n5["sequence"].tolist()
    y_h5n5 = df_h5n5["label"].values.astype(int)

    X_h5n5, artifacts = build_kmer_features(
        sequences=sequences,
        k=CFG.kmer_k,
        svd_components=CFG.svd_components
    )

    dump(artifacts, "data/processed/feature_artifacts.joblib")

    X_train_h5n5, X_test_h5n5, y_train_h5n5, y_test_h5n5 = train_test_split(
        X_h5n5, y_h5n5,
        test_size=CFG.test_size,
        random_state=CFG.random_state,
        stratify=y_h5n5
    )

    print("\nTraining Classical SVM (RBF) with grid search...")
    h5n5_classical = train_classical_svm(X_train_h5n5, y_train_h5n5)
    dump(h5n5_classical.model, "data/processed/classical_svm.joblib")
    print("Best CV:", h5n5_classical.best_cv_score, "Params:", h5n5_classical.best_params)
    acc_c_h5n5 = evaluate(h5n5_classical.model, X_test_h5n5, y_test_h5n5, "Classical SVM (H5N5)")

    print("\nTraining Quantum SVM (QSVC, quantum kernel)...")
    h5n5_quantum = train_quantum_svm(X_train_h5n5, y_train_h5n5, n_qubits=CFG.n_qubits)
    dump(h5n5_quantum.model, "data/processed/quantum_qsvc.joblib")
    acc_q_h5n5 = evaluate(h5n5_quantum.model, X_test_h5n5, y_test_h5n5, "Quantum SVM (H5N5)")

    print("\nTraining Naive Bayes...")
    h5n5_nb = train_naive_bayes(X_train_h5n5, y_train_h5n5)
    dump(h5n5_nb.model, "data/processed/h5n5_nb.joblib")
    acc_nb_h5n5 = evaluate(h5n5_nb.model, X_test_h5n5, y_test_h5n5, "Naive Bayes (H5N5)")

    print("\nTraining Decision Tree (Grid Search)...")
    h5n5_dt = train_decision_tree(X_train_h5n5, y_train_h5n5)
    dump(h5n5_dt.model, "data/processed/h5n5_dt.joblib")
    acc_dt_h5n5 = evaluate(h5n5_dt.model, X_test_h5n5, y_test_h5n5, "Decision Tree (H5N5)")

    h5n5_results = {
        "n_samples": int(len(df_h5n5)),
        "n_pos": int(df_h5n5["label"].sum()),
        "n_neg": int((df_h5n5["label"] == 0).sum()),
        "kmer_k": CFG.kmer_k,
        "svd_components": CFG.svd_components,
        "classical_acc": float(acc_c_h5n5),
        "quantum_acc": float(acc_q_h5n5),
        "rf_acc": float(acc_rf_h5n5),
        "nb_acc": float(acc_nb_h5n5),
        "dt_acc": float(acc_dt_h5n5),
        "classical_best_params": h5n5_classical.best_params,
        "rf_best_params": h5n5_rf.best_params,
        "dt_best_params": h5n5_dt.best_params
    }

    # ==================== PIMA ANALYSIS ====================
    print("\n" + "="*70)
    print("PIMA DIABETES ANALYSIS - EDA")
    print("="*70)
    
    pima_eda = PimaEDA(pd.DataFrame())
    df_pima = pima_eda.load_pima_data()
    pima_eda.df = df_pima
    
    eda_stats = pima_eda.basic_statistics()
    print(f"Dataset shape: {eda_stats['shape']}")
    print("Generating visualizations...")
    
    pima_eda.distribution_plots()
    pima_eda.correlation_heatmap()
    pima_eda.class_distribution()
    pima_eda.feature_vs_target()
    
    # Train models on Pima
    X_pima = df_pima.drop('Outcome', axis=1).values
    y_pima = df_pima['Outcome'].values
    
    X_train_pima, X_test_pima, y_train_pima, y_test_pima = train_test_split(
        X_pima, y_pima,
        test_size=CFG.test_size,
        random_state=CFG.random_state,
        stratify=y_pima
    )
    
    acc_pima_c, acc_pima_q, acc_pima_rf, acc_pima_nb, acc_pima_dt, X_test_pima_scaled, X_test_pima_q, y_test_pima, pima_c_model, pima_q_model, pima_rf_model, pima_nb_model, pima_dt_model = train_pima_models(
        X_train_pima, X_test_pima, y_train_pima, y_test_pima
    )
    
    pima_results = {
        "classical_acc": float(acc_pima_c),
        "quantum_acc": float(acc_pima_q),
        "rf_acc": float(acc_pima_rf),
        "nb_acc": float(acc_pima_nb),
        "dt_acc": float(acc_pima_dt),
    }
    
    # ==================== VISUALIZATIONS ====================
    print("\n" + "="*70)
    print("GENERATING VISUALIZATIONS")
    print("="*70)
    
    viz = ResultsVisualizer()
    viz.plot_model_comparison(acc_c_h5n5, acc_q_h5n5, acc_rf_h5n5, acc_nb_h5n5, acc_dt_h5n5, 
                             acc_pima_c, acc_pima_q, acc_pima_rf, acc_pima_nb, acc_pima_dt)
    
    # Confusion matrices
    viz.plot_confusion_matrices(
        y_test_h5n5,
        {
            "Classical SVM": h5n5_classical.model.predict(X_test_h5n5),
            "Quantum SVM": h5n5_quantum.model.predict(X_test_h5n5),
            "Random Forest": h5n5_rf.model.predict(X_test_h5n5),
            "Naive Bayes": h5n5_nb.model.predict(X_test_h5n5),
            "Decision Tree": h5n5_dt.model.predict(X_test_h5n5)
        },
        "H5N5"
    )
    
    viz.plot_confusion_matrices(
        y_test_pima,
        {
            "Classical SVM": pima_c_model.model.predict(X_test_pima_scaled),
            "Quantum SVM": pima_q_model.model.predict(X_test_pima_q),
            "Random Forest": pima_rf_model.model.predict(X_test_pima_scaled),
            "Naive Bayes": pima_nb_model.model.predict(X_test_pima_scaled),
            "Decision Tree": pima_dt_model.model.predict(X_test_pima_scaled)
        },
        "Pima"
    )
    
    # ==================== REPORTING ====================
    print("\n" + "="*70)
    print("GENERATING WORD REPORT")
    print("="*70)
    
    generate_word_report(h5n5_results, pima_results, eda_stats)
    
    # Save summary JSON
    with open("data/processed/run_summary.json", "w") as f:
        json.dump({
            "h5n5": h5n5_results,
            "pima": pima_results,
            "timestamp": datetime.now().isoformat()
        }, f, indent=2)
    print("✓ Summary saved: data/processed/run_summary.json")
    
    print("\n" + "="*70)
    print("ANALYSIS COMPLETE")
    print("="*70)
    print("Outputs:")
    print("  - H5N5 Dataset: data/processed/h5n5_dataset.csv")
    print("  - Pima Dataset: data/raw/pima_indians_diabetes.csv")
    print("  - Visualizations: data/processed/*.png")
    print("  - Word Report: data/processed/ML_Classification_Report.docx")
    print("  - Summary: data/processed/run_summary.json")


if __name__ == "__main__":
    main()
